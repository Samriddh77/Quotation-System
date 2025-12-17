import streamlit as st
import pandas as pd
import os
import re
from datetime import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    st.error("❌ Library 'python-docx' is missing. Please run: pip install python-docx")
    st.stop()

# --- 1. SETUP PATHS ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, 'templates')
DATA_DIR = os.path.join(BASE_DIR, 'data')

# Map Firm Names to Template Files
FIRM_MAPPING = {
    "Electro World": "Electro_Template.docx",
    "Abhinav Enterprises": "Abhinav_Template.docx",
    "Shree Creative Marketing": "Shree_Template.docx"
}

# --- 2. FOLDER SETUP ---
if not os.path.exists(DATA_DIR): os.makedirs(DATA_DIR)
if not os.path.exists(TEMPLATE_DIR): os.makedirs(TEMPLATE_DIR)

# --- CONFIGURATION ---
st.set_page_config(page_title="Quotation Generator", layout="wide")

# --- AUTO-FILL DEFAULTS ---
FIRM_DEFAULTS = {
    "Electro World": {
        "prefix": "EW/QTN", "price": "Nett Ex-Works", "gst": "Extra @ 18% as applicable",
        "delivery": "Ex-Stock / 1-2 Weeks", "freight": "To Pay / Ex-Godown",
        "payment": "100% Against Proforma Invoice", "validity": "7 Days", "guarantee": "12 Months"
    },
    "Abhinav Enterprises": {
        "prefix": "AE/QTN", "price": "Nett F.O.R.", "gst": "18% Extra",
        "delivery": "Within 2-3 Weeks", "freight": "P&F: NIL | F.O.R: Ex Indore",
        "payment": "30% Advance", "validity": "15 Days", "guarantee": "12 Months"
    },
    "Shree Creative Marketing": {
        "prefix": "SCM/QTN", "price": "Nett", "gst": "Inclusive",
        "delivery": "Immediate", "freight": "Paid by us",
        "payment": "COD", "validity": "5 Days", "guarantee": "Standard"
    }
}

# --- STATE MANAGEMENT ---
def update_defaults():
    firm = st.session_state.get('firm_selector', "Electro World")
    defaults = FIRM_DEFAULTS.get(firm, FIRM_DEFAULTS["Electro World"])
    st.session_state['p_term'] = defaults['price']
    st.session_state['g_term'] = defaults['gst']
    st.session_state['d_term'] = defaults['delivery']
    st.session_state['f_term'] = defaults['freight']
    st.session_state['pay_term'] = defaults['payment']
    st.session_state['val_term'] = defaults['validity']
    st.session_state['guar_term'] = defaults['guarantee']
    date_str = datetime.now().strftime('%y%m%d')
    st.session_state['ref_no_val'] = f"{defaults['prefix']}/{date_str}/001"

if 'firm_selector' not in st.session_state: st.session_state['firm_selector'] = "Electro World"
if 'p_term' not in st.session_state: update_defaults()

# --- HELPERS ---
def clean_price_value(val):
    if pd.isna(val): return 0.0
    s = str(val).strip()
    s_clean = re.sub(r'[^\d.]', '', s)
    try: return float(s_clean)
    except: return 0.0

def clean_coil_len(val):
    if pd.isna(val): return 0.0
    s = str(val).strip()
    s_clean = re.sub(r'[^\d.]', '', s)
    try: return float(s_clean)
    except: return 0.0

def detect_uom(sheet_name, price_col_name):
    s_up = str(sheet_name).upper()
    c_up = str(price_col_name).upper()
    if "MTR" in c_up or "METER" in c_up: return "Mtr"
    if "PC" in c_up or "PIECE" in c_up: return "Pc"
    if "GLAND" in s_up or "HMI" in s_up or "COSMOS" in s_up: return "Pc"
    return "Mtr" 

def format_qty(qty, uom):
    uom_clean = str(uom).lower().strip()
    if any(x == uom_clean for x in ['pc', 'no', 'nos', 'set', 'each', 'fix']):
        return f"{int(qty)}"
    return f"{qty:,.2f}"

# --- WORD GENERATOR ---
def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblCellMar = tblPr.first_child_found_in("w:tblCellMar")
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar')
        tblPr.append(tblCellMar)
    for side, w in [("top", "60"), ("bottom", "60"), ("left", "100"), ("right", "100")]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), w)
        node.set(qn("w:type"), "dxa")
        tblCellMar.append(node)

def safe_replace_text(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text: run.text = run.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text: run.text = run.text.replace(key, value)

def fill_template_docx(template_path, client_data, cart_items, terms, visible_cols):
    doc = Document(template_path)
    
    replacements = {
        '{{REF_NO}}': str(client_data.get('ref_no', '')),
        '{{DATE}}': datetime.now().strftime('%d-%b-%Y'),
        '{{CLIENT_NAME}}': str(client_data.get('client_name', '')),
        '{{CLIENT_ADDRESS}}': str(client_data.get('client_address', '')),
        '{{SUBJECT}}': str(client_data.get('subject', '')),
        '{{PRICE_TERM}}': str(terms.get('price_term', '')),
        '{{GST_TERM}}': str(terms.get('gst_term', '')),
        '{{DELIVERY_TERM}}': str(terms.get('delivery_term', '')),
        '{{FREIGHT_TERM}}': str(terms.get('freight_term', '')),
        '{{PAYMENT_TERM}}': str(terms.get('payment_term', '')),
        '{{VALIDITY_TERM}}': str(terms.get('validity_term', '')),
        '{{GUARANTEE_TERM}}': str(terms.get('guarantee_term', '')),
        '{{GURANTEE_TERM}}': str(terms.get('guarantee_term', '')),
    }

    # Replace in Paragraphs & Tables
    safe_replace_text(doc, replacements)

    # Insert Table at {{TABLE_HERE}}
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if '{{TABLE_HERE}}' in paragraph.text:
            target_paragraph = paragraph
            break
            
    if target_paragraph:
        target_paragraph.text = "" 
        
        # Column Ratios
        col_ratios = {"S.No.": 5, "Sub Category": 12, "Item Description": 35, "Make": 10, "Qty": 8, "Unit": 8, "Rate": 10, "Amount": 12}
        active_headers = [h for h in visible_cols if h in col_ratios]
        
        table = doc.add_table(rows=1, cols=len(active_headers))
        table.autofit = False
        table.allow_autofit = False
        
        # Set Table Width to 100%
        tblPr = table._tbl.tblPr
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000') 
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)
        
        set_table_borders(table)
        
        # Header
        for i, text in enumerate(active_headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        total_amt = 0
        for i, item in enumerate(cart_items):
            row_cells = table.add_row().cells
            lp = item['List Price']
            disc = item['Discount']
            qty = item['Qty']
            net_rate = lp * (1 - disc/100)
            line_total = net_rate * qty
            total_amt += line_total
            
            desc = item['Description']
            make_str = f"{item['Make']} Make" if item['Make'].strip() else ""
            qty_fmt = format_qty(qty, item['Display Unit'])
            
            data_map = {
                "S.No.": str(i+1), "Sub Category": item.get('Sub Category', ''),
                "Item Description": desc, "Make": make_str,
                "Qty": qty_fmt, "Unit": item['Display Unit'],
                "Rate": f"{net_rate:,.2f}", "Amount": f"{line_total:,.2f}"
            }
            
            for idx, header in enumerate(active_headers):
                cell = row_cells[idx]
                cell.text = data_map.get(header, "")
                if header in ["Qty", "Rate", "Amount"]: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif header in ["S.No.", "Unit", "Make"]: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Total Row
        if "Amount" in active_headers:
            row = table.add_row().cells
            amt_idx = active_headers.index("Amount")
            label_idx = max(0, amt_idx - 1)
            
            row[label_idx].text = "Grand Total (Excl. GST)"
            row[label_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if not row[label_idx].paragraphs[0].runs: row[label_idx].paragraphs[0].add_run("Grand Total (Excl. GST)").bold = True
            else: row[label_idx].paragraphs[0].runs[0].bold = True

            row[amt_idx].text = f"{total_amt:,.2f}"
            row[amt_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if not row[amt_idx].paragraphs[0].runs: row[amt_idx].paragraphs[0].add_run(f"{total_amt:,.2f}").bold = True
            else: row[amt_idx].paragraphs[0].runs[0].bold = True

        target_paragraph._p.addnext(table._tbl)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- DATA LOADER ---
@st.cache_data(show_spinner=True)
def load_data_from_files():
    # Use Dynamic Path for Data Directory
    if not os.path.exists(DATA_DIR):
        return pd.DataFrame(), [f"❌ Data folder not found at: {DATA_DIR}"]

    all_dfs = []
    logs = []
    
    excel_files = [os.path.join(DATA_DIR, f) for f in os.listdir(DATA_DIR) 
                   if f.lower().endswith(".xlsx") and not f.startswith("~$")]
    
    if not excel_files: return pd.DataFrame(), ["❌ No .xlsx files found in data folder!"]

    for file_path in excel_files:
        filename = os.path.basename(file_path)
        main_cat = os.path.splitext(filename)[0] 
        try:
            xls = pd.ExcelFile(file_path)
            for sheet in xls.sheet_names:
                try:
                    df = pd.read_excel(xls, sheet)
                    df.columns = [str(c).strip() for c in df.columns]
                    
                    # 1. Identify Columns
                    name_col = next((c for c in df.columns if "Item Description" == c), None)
                    price_col = next((c for c in df.columns if "List Price" == c), None)
                    disc_col = next((c for c in df.columns if "Standard Discount" == c), None)
                    coil_col = next((c for c in df.columns if "Coil Length (Mtr)" == c), None)
                    uom_col = next((c for c in df.columns if "UOM" == c), None)

                    if name_col and price_col:
                        clean_df = pd.DataFrame()
                        clean_df['Description'] = df[name_col].astype(str)
                        clean_df['List Price'] = df[price_col].apply(clean_price_value)
                        
                        # Load Discount
                        if disc_col:
                            clean_df['Standard Discount'] = pd.to_numeric(df[disc_col], errors='coerce').fillna(0)
                        else:
                            clean_df['Standard Discount'] = 0.0

                        # Load Coil Length
                        if coil_col:
                            clean_df['Coil Length'] = df[coil_col].apply(clean_coil_len)
                        else:
                            clean_df['Coil Length'] = 0.0

                        clean_df['Main Category'] = main_cat
                        clean_df['Sub Category'] = sheet
                        
                        # Load UOM
                        if uom_col:
                            clean_df['UOM'] = df[uom_col].astype(str)
                        else:
                            clean_df['UOM'] = detect_uom(sheet, price_col)

                        clean_df = clean_df[clean_df['List Price'] > 0]
                        all_dfs.append(clean_df)
                except: pass
        except: pass

    if not all_dfs: return pd.DataFrame(), logs
    return pd.concat(all_dfs, ignore_index=True), logs

# --- APP UI ---
catalog, logs = load_data_from_files()
if 'cart' not in st.session_state: st.session_state['cart'] = []

# SIDEBAR
with st.sidebar:
    st.title("🔧 Config")
    if st.button("🔄 Refresh Data"): st.cache_data.clear(); st.rerun()

    st.markdown("---")
    st.header("1. Add Item")
    if not catalog.empty:
        cats = sorted(catalog['Main Category'].unique())
        sel_cat = st.selectbox("Category", cats)
        sub_cats = sorted(catalog[catalog['Main Category'] == sel_cat]['Sub Category'].unique())
        sel_sub = st.selectbox("Sub Category", sub_cats)
        
        subset = catalog[(catalog['Main Category'] == sel_cat) & (catalog['Sub Category'] == sel_sub)]
        prods = sorted(subset['Description'].unique())
        sel_prod = st.selectbox("Product", prods)
        
        # Fetch Data for Selected Product
        row = subset[subset['Description'] == sel_prod].iloc[0]
        std_price = row['List Price']
        std_disc = row['Standard Discount']
        uom_raw = row['UOM']
        coil_len = row['Coil Length']
        
        st.info(f"Price: {std_price} / {uom_raw} | Discount: {std_disc}%")
        if coil_len > 0:
            st.caption(f"ℹ️ Standard Coil Length: {coil_len} {uom_raw}")
        
        # --- QUANTITY LOGIC ---
        calc_qty = 0
        disp_unit = uom_raw
        
        # If Coil Length exists and UOM is Meter (e.g. Wires), show toggle
        if coil_len > 0 and "MTR" in str(uom_raw).upper():
            mode = st.radio("Input Mode", ["Coils", "Meters"], horizontal=True)
            if mode == "Coils":
                n = st.number_input("No. of Coils", 1, 500)
                calc_qty = n * coil_len
                # Keep unit as Mtr for price calc, but maybe note coils in cart
                disp_unit = "Mtr" 
            else:
                calc_qty = st.number_input("Total Meters", 1.0, 10000.0)
                disp_unit = "Mtr"
        else:
            # Standard Items (Glands, Cables without fixed coil len)
            if any(x in str(uom_raw).upper() for x in ["PC", "NO", "SET"]):
                calc_qty = st.number_input(f"Qty ({uom_raw})", 1, 10000)
            else:
                calc_qty = st.number_input(f"Qty ({uom_raw})", 1.0, 10000.0)

        c1, c2 = st.columns(2)
        # Pre-fill discount with value from Excel
        disc = c1.number_input("Disc %", 0.0, 100.0, float(std_disc))
        make = c2.text_input("Make")
        
        if st.button("Add"):
            if not make.strip():
                st.error("⚠️ The 'Make' field is mandatory!")
            else:
                st.session_state['cart'].append({
                    'Main Category': sel_cat, 'Sub Category': sel_sub,
                    'Description': sel_prod, 'Make': make,
                    'Qty': calc_qty, 'Display Unit': disp_unit,
                    'List Price': std_price, 'Discount': disc,
                    'Sub Category': sel_sub # Ensure this is saved
                })
                st.success("Added")
    else:
        st.warning("Please add an Excel file to the 'data' folder.")
            
    if st.button("Clear Cart"): st.session_state['cart'] = []; st.rerun()

# MAIN PAGE
st.title("📄 Quotation System")

if not st.session_state['cart']:
    st.info("Add items to start.")
else:
    # TABLE
    st.subheader("1. Item List")
    col_config = [0.5, 3.5, 1.5, 1.5, 1.2, 1.5, 0.5]
    h1, h2, h3, h4, h5, h6, h7 = st.columns(col_config)
    h1.write("#"); h2.write("Desc"); h3.write("Make"); h4.write("Qty"); h5.write("Unit"); h6.write("Total"); 
    st.divider()

    grand_tot = 0
    for i, item in enumerate(st.session_state['cart']):
        net = item['List Price'] * (1 - item['Discount']/100)
        tot = net * item['Qty']
        grand_tot += tot
        
        c1, c2, c3, c4, c5, c6, c7 = st.columns(col_config)
        c1.write(f"{i+1}")
        c2.write(item['Description'])
        c3.write(f"{item['Make']} Make")
        
        # Display Format
        qty_str = format_qty(item['Qty'], item['Display Unit'])
        c4.write(qty_str)
        c5.write(item['Display Unit'])
        c6.write(f"₹ {tot:,.0f}")
        if c7.button("🗑️", key=f"d{i}"):
            st.session_state['cart'].pop(i)
            st.rerun()

    st.divider()
    st.write(f"**Est. Grand Total: ₹ {grand_tot:,.2f}**")
    st.markdown("---")

    # GENERATOR FORM
    st.subheader("2. Generate Official Quotation")
    
    selected_firm = st.selectbox("Select Template (Firm)", list(FIRM_MAPPING.keys()), key="firm_selector", on_change=update_defaults)
    
    st.markdown("##### Table Settings")
    all_cols = ["S.No.", "Sub Category", "Item Description", "Make", "Qty", "Unit", "Rate", "Amount"]
    visible_cols = st.multiselect("Columns to include", all_cols, default=all_cols)

    c1, c2 = st.columns(2)
    with c1:
        client_name = st.text_input("Client Name", placeholder="M/s Client Name")
        ref_no = st.text_input("Ref No", key="ref_no_val")
        subject = st.text_input("Subject", value="OFFER FOR CABLES / ELECTRICAL GOODS")
    with c2:
        client_address = st.text_area("Client Address", placeholder="Address Line 1\nCity, State")
    
    st.markdown("#### Terms & Conditions (Editable)")
    tc1, tc2, tc3 = st.columns(3)
    p_term = tc1.text_input("Price Term", key='p_term')
    g_term = tc2.text_input("GST Term", key='g_term')
    d_term = tc3.text_input("Delivery", key='d_term')
    
    tc4, tc5, tc6 = st.columns(3)
    f_term = tc4.text_input("Freight", key='f_term')
    pay_term = tc5.text_input("Payment", key='pay_term')
    val_term = tc6.text_input("Validity", key='val_term')
    guarantee = st.text_input("Guarantee", key='guar_term')

    st.markdown("---")
    st.markdown("### 3. Output Options")
    
    # Generate the filename first
    safe_client = "".join([c for c in client_name if c.isalnum() or c in (' ', '_')]).strip()
    filename = f"Quote_{selected_firm.replace(' ','_')}_{safe_client[:10]}_{datetime.now().strftime('%d%b')}.docx"

    # --- OPTION A: BROWSER DOWNLOAD (Standard) ---
    st.write("#### Option A: Download via Browser")
    if st.button("📥 Generate & Download", type="primary"):
        if not client_name:
            st.error("Please enter Client Name.")
        else:
            # Prepare Data
            client_data = {
                "client_name": client_name, "client_address": client_address,
                "ref_no": ref_no, "subject": subject
            }
            terms = {
                "price_term": p_term, "gst_term": g_term, "delivery_term": d_term,
                "freight_term": f_term, "payment_term": pay_term, "validity_term": val_term,
                "guarantee_term": guarantee
            }
            
            # Generate Logic
            template_path = os.path.join(TEMPLATE_DIR, FIRM_MAPPING[selected_firm])
            if not os.path.exists(template_path):
                 st.error(f"❌ Template not found: {template_path}")
            else:
                docx_buffer = fill_template_docx(template_path, client_data, st.session_state['cart'], terms, visible_cols)
                
                # Trigger Browser Download
                st.download_button(
                    label="⬇️ Click here to Download",
                    data=docx_buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # --- OPTION B: SAVE TO SPECIFIC LOCAL FOLDER ---
    st.write("#### Option B: Save to Specific Local Folder")
    
    # 1. Input for Custom Path (Default to a folder named 'Generated_Quotes' in your project dir)
    default_save_path = os.path.join(BASE_DIR, 'Generated_Quotes')
    save_folder = st.text_input("Target Folder Path", value=default_save_path)

    if st.button("💾 Save to Local Folder"):
        if not client_name:
            st.error("Please enter Client Name.")
        else:
            # Check if folder exists, if not, create it
            if not os.path.exists(save_folder):
                try:
                    os.makedirs(save_folder)
                    st.toast(f"Created new folder: {save_folder}")
                except Exception as e:
                    st.error(f"Could not create folder: {e}")
                    st.stop()

            # Prepare Data (Same as above)
            client_data = {
                "client_name": client_name, "client_address": client_address,
                "ref_no": ref_no, "subject": subject
            }
            terms = {
                "price_term": p_term, "gst_term": g_term, "delivery_term": d_term,
                "freight_term": f_term, "payment_term": pay_term, "validity_term": val_term,
                "guarantee_term": guarantee
            }

            # Generate Logic
            template_path = os.path.join(TEMPLATE_DIR, FIRM_MAPPING[selected_firm])
            if not os.path.exists(template_path):
                 st.error(f"❌ Template not found: {template_path}")
            else:
                docx_buffer = fill_template_docx(template_path, client_data, st.session_state['cart'], terms, visible_cols)
                
                full_save_path = os.path.join(save_folder, filename)
                
                try:
                    with open(full_save_path, "wb") as f:
                        f.write(docx_buffer.getbuffer())
                    st.success(f"✅ Successfully saved to: {full_save_path}")
                except Exception as e:
                    st.error(f"❌ Error saving file: {e}")